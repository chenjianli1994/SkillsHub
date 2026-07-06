#!/usr/bin/env python3
"""Generate a strategy document DOCX from structured JSON and a DOCX template."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, list):
        return "\n".join(text(item) for item in value if text(item))
    return str(value)


def remove_all_children(element) -> None:
    for child in list(element):
        element.remove(child)


def clear_document_body(document: Document) -> None:
    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    if sect_pr is not None and (len(body) == 0 or body[-1] is not sect_pr):
        if sect_pr.getparent() is not None:
            sect_pr.getparent().remove(sect_pr)
        body.append(sect_pr)


def set_page_layout(document: Document) -> None:
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.4)
        section.top_margin = Cm(2.08)
        section.bottom_margin = Cm(2.5)
        section.header_distance = Cm(0.0)
        section.footer_distance = Cm(1.7)
        section.start_type = WD_SECTION_START.NEW_PAGE


def style_name(document: Document, preferred: str, fallback: str = "Normal") -> str:
    names = {style.name for style in document.styles}
    if preferred in names:
        return preferred
    return fallback


def table_style_name(document: Document) -> str | None:
    for candidate in ["Table Grid", "Normal Table", "TableNormal"]:
        if candidate in {style.name for style in document.styles}:
            return candidate
    return None


def set_run_font(run, east_asia: str = "宋体", size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def apply_cell_shading(cell, fill: str = "EDEDED") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, value: Any, *, bold: bool = False, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text(value))
    set_run_font(run, size_pt=10.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_field_run(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def set_header_footer(document: Document, title: str, version: str, status: str) -> None:
    for section in document.sections:
        header = section.header
        remove_all_children(header._element)
        p = header.add_paragraph(style=style_name(document, "Header Footer"))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        run = p.add_run(f"{title}    {version}    {status}")
        set_run_font(run, size_pt=9)

        footer = section.footer
        remove_all_children(footer._element)
        p = footer.add_paragraph(style=style_name(document, "Header Footer"))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("第 ")
        set_run_font(run, size_pt=9)
        add_field_run(p, "PAGE")
        run = p.add_run(" / ")
        set_run_font(run, size_pt=9)
        add_field_run(p, "NUMPAGES")
        run = p.add_run(" 页")
        set_run_font(run, size_pt=9)


def add_toc(document: Document) -> None:
    document.add_paragraph("目录", style=style_name(document, "TOC Heading", "Heading 1"))
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:r")
    hint_text = OxmlElement("w:t")
    hint_text.text = "打开 Word/WPS 后更新域，可生成带页码的目录。"
    hint.append(hint_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    paragraph._p.append(hint)
    run._r.append(fld_end)


def add_cover(document: Document, data: dict[str, Any]) -> None:
    info = data.get("document_info", {})
    approval = data.get("approval_info", {})

    title = text(info.get("文档标题"), "策略文档")
    subtitle = text(info.get("文档副标题"), "控制策略")
    version = text(info.get("版本"), "r01")
    status = text(info.get("状态"), "待评审")
    unit = text(approval.get("单位") or info.get("编写单位"))
    date = text(approval.get("日期") or info.get("编写日期"))

    p = document.add_paragraph("策略文档", style=style_name(document, "Cover Kicker"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(title, style=style_name(document, "Cover Title"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"{subtitle} / {version} / {status}", style=style_name(document, "Cover Subtitle"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")

    table = document.add_table(rows=5, cols=2)
    chosen_style = table_style_name(document)
    if chosen_style:
        table.style = chosen_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["编制", "校核", "会签", "批准", "来源说明"]
    values = [
        approval.get("编制", ""),
        approval.get("校核", ""),
        approval.get("会签", ""),
        approval.get("批准", ""),
        "需求文档 + 代码实现",
    ]
    for row_idx, (label, value) in enumerate(zip(labels, values)):
        set_cell_text(table.rows[row_idx].cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        apply_cell_shading(table.rows[row_idx].cells[0], "F3F3F3")
        set_cell_text(table.rows[row_idx].cells[1], value)

    document.add_paragraph("")
    if unit:
        p = document.add_paragraph(unit, style=style_name(document, "Caption", "Normal"))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if date:
        p = document.add_paragraph(date, style=style_name(document, "Caption", "Normal"))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


def add_revision_history(document: Document, revisions: list[dict[str, Any]]) -> None:
    document.add_paragraph("修订记录", style=style_name(document, "Heading 1"))
    table = document.add_table(rows=1, cols=5)
    chosen_style = table_style_name(document)
    if chosen_style:
        table.style = chosen_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["NO", "修订内容", "修订人", "版本号", "发布日期"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        apply_cell_shading(table.rows[0].cells[idx])
    for item in revisions or []:
        row = table.add_row().cells
        values = [
            item.get("NO", ""),
            item.get("修订内容", ""),
            item.get("修订人", ""),
            item.get("版本号", ""),
            item.get("发布日期", ""),
        ]
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)
    document.add_page_break()


def set_column_widths(table, widths_cm: list[float]) -> None:
    if not widths_cm:
        return
    for row in table.rows:
        for idx, width in enumerate(widths_cm[: len(row.cells)]):
            row.cells[idx].width = Cm(width)


def add_table(document: Document, block: dict[str, Any]) -> None:
    title = text(block.get("title"))
    if title:
        document.add_paragraph(title, style=style_name(document, "Caption", "Normal"))
    headers = [text(item) for item in block.get("headers", [])]
    rows = block.get("rows", [])
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    chosen_style = table_style_name(document)
    if chosen_style:
        table.style = chosen_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(table, block.get("column_widths_cm", []))
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        apply_cell_shading(table.rows[0].cells[idx])
    for item in rows:
        row = table.add_row().cells
        values = item if isinstance(item, list) else [item.get(header, "") for header in headers]
        for idx, value in enumerate(values[: len(headers)]):
            set_cell_text(row[idx], value)
    document.add_paragraph("")


def add_image(document: Document, block: dict[str, Any]) -> None:
    path = Path(text(block.get("path")))
    if not path.exists():
        note = document.add_paragraph(style=style_name(document, "Caption", "Normal"))
        note.add_run(f"图片未找到：{path}")
        return
    width_cm = float(block.get("width_cm", 14))
    document.add_picture(str(path), width=Cm(width_cm))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = text(block.get("caption"))
    if caption:
        p = document.add_paragraph(caption, style=style_name(document, "Caption", "Normal"))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_list(document: Document, values: list[Any], ordered: bool) -> None:
    for item in values:
        p = document.add_paragraph(style=style_name(document, "List Paragraph", "Normal"))
        p.style = document.styles[style_name(document, "List Paragraph", "Normal")]
        prefix = "" if ordered else "• "
        run = p.add_run(prefix + text(item))
        set_run_font(run, size_pt=10.5)


def add_paragraph_block(document: Document, value: str, style: str = "Normal") -> None:
    p = document.add_paragraph(value, style=style_name(document, style, "Normal"))
    for run in p.runs:
        set_run_font(run, size_pt=10.5 if style == "Normal" else None)


def add_heading_block(document: Document, title: str, level: int) -> None:
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    document.add_paragraph(title, style=style_name(document, style_map.get(level, "Heading 3"), "Heading 3"))


def render_block(document: Document, block: dict[str, Any]) -> None:
    block_type = block.get("type")
    if block_type == "heading":
        add_heading_block(document, text(block.get("title")), int(block.get("level", 2)))
    elif block_type == "paragraph":
        add_paragraph_block(document, text(block.get("text")))
    elif block_type == "note":
        add_paragraph_block(document, text(block.get("text")), style="Caption")
    elif block_type == "bullet_list":
        add_list(document, block.get("items", []), ordered=False)
    elif block_type == "numbered_list":
        add_list(document, block.get("items", []), ordered=True)
    elif block_type == "table":
        add_table(document, block)
    elif block_type == "image":
        add_image(document, block)
    elif block_type == "page_break":
        document.add_page_break()
    else:
        add_paragraph_block(document, text(block))


def render_sections(document: Document, sections: list[dict[str, Any]]) -> None:
    for index, section in enumerate(sections):
        add_heading_block(document, text(section.get("title"), f"第{index + 1}节"), 1)
        for block in section.get("blocks", []):
            render_block(document, block)


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8-sig") as handle:
        return json.load(handle)


def build_document(data: dict[str, Any], template: Path, output: Path) -> None:
    document = Document(str(template))
    clear_document_body(document)
    set_page_layout(document)

    info = data.get("document_info", {})
    set_header_footer(
        document,
        text(info.get("文档标题"), "策略文档"),
        text(info.get("版本"), "r01"),
        text(info.get("状态"), "待评审"),
    )

    add_cover(document, data)
    add_revision_history(document, data.get("revision_history", []))
    add_toc(document)
    document.add_page_break()
    render_sections(document, data.get("sections", []))
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="按 Word 模板生成策略文档 DOCX")
    parser.add_argument("input", help="策略文档结构化 JSON 文件")
    parser.add_argument("-t", "--template", required=True, help="Word 模板 DOCX 文件")
    parser.add_argument("-o", "--output", required=True, help="输出 DOCX 文件")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    build_document(load_json(Path(args.input)), Path(args.template), Path(args.output))
    print(Path(args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
